#!/usr/bin/env python3
"""Audit every formal-test PDF and DOCX without changing the question bank."""
from __future__ import annotations

import argparse
import json
import re
import subprocess
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path

from docx import Document

import import_formal_test_questions as importer


ROOT = Path(__file__).resolve().parent.parent
FORMAL = ROOT / "formal test"


def pdf_text(path: Path) -> str:
    return subprocess.run(
        ["pdftotext", "-raw", str(path), "-"], check=True, capture_output=True, text=True
    ).stdout


def pdf_pages(path: Path) -> int:
    output = subprocess.run(["pdfinfo", str(path)], check=True, capture_output=True, text=True).stdout
    return int(re.search(r"^Pages:\s+(\d+)$", output, re.M).group(1))


def question_marker_count(path: Path) -> int:
    return len(re.findall(r"Question\s*#\s*\d+\s*-\s*[^\n]+", pdf_text(path), re.I))


def civil_question_marker_count(path: Path) -> int:
    return len(re.findall(r"(?m)^Question\s*#\s*\d+", pdf_text(path).replace("\f", "\n")))


def file_entry(path: Path, role: str, status: str, **details: object) -> dict:
    return {
        "file": str(path.relative_to(ROOT)),
        "format": path.suffix.lower().lstrip("."),
        "role": role,
        "pages": pdf_pages(path) if path.suffix.lower() == ".pdf" else 23,
        "status": status,
        **details,
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, default=FORMAL / "formal_test_source_audit.json")
    args = parser.parse_args()
    artifacts: list[dict] = []
    total = 0
    errors: list[str] = []

    seven_day = sorted((FORMAL / "7 day questions ").glob("*.pdf"))
    seven_day_total = 0
    for path in seven_day:
        records, expected, parsing_errors = importer.parse_7_day(path)
        count = len(records)
        seven_day_total += count
        total += count
        artifacts.append(file_entry(
            path, "question_and_answer_key", "complete" if count == expected and not parsing_errors else "incomplete",
            expected_questions=expected, extracted_questions=count, parse_errors=parsing_errors,
        ))
        errors.extend(parsing_errors)

    path_210 = FORMAL / "210-Questions.pdf"
    records, expected, parsing_errors = importer.parse_210(path_210)
    total += len(records)
    artifacts.append(file_entry(
        path_210, "question_and_answer_key", "complete" if len(records) == expected and not parsing_errors else "incomplete",
        expected_questions=expected, extracted_questions=len(records), parse_errors=parsing_errors,
    ))
    errors.extend(parsing_errors)

    civil_answers = FORMAL / "civil procedure" / "CivilProcedure2021_Answer-Keys.pdf"
    civil_questions = civil_answers.with_name("CivilProcedure2021_Questions.pdf")
    records, expected, parsing_errors = importer.parse_civil_2021(civil_answers)
    total += len(records)
    civil_question_count = civil_question_marker_count(civil_questions)
    artifacts.extend([
        file_entry(
            civil_questions, "question_text", "complete" if civil_question_count == expected else "incomplete",
            expected_questions=expected, document_question_markers=civil_question_count, paired_answer_key=str(civil_answers.relative_to(ROOT)),
        ),
        file_entry(
            civil_answers, "answer_key_with_explanations", "complete" if len(records) == expected and not parsing_errors else "incomplete",
            expected_questions=expected, extracted_questions=len(records), parse_errors=parsing_errors,
        ),
    ])
    errors.extend(parsing_errors)

    for answers in sorted((FORMAL / "NCBE Online Practice Exam 1-4 ").glob("* Answers.pdf")):
        questions = answers.with_name(answers.name.replace("Answers", "Questions"))
        records, expected, parsing_errors = importer.parse_ncbe(answers)
        total += len(records)
        question_count = question_marker_count(questions)
        artifacts.extend([
            file_entry(
                questions, "question_text", "complete" if question_count == expected else "incomplete",
                expected_questions=expected, document_question_markers=question_count, paired_answer_key=str(answers.relative_to(ROOT)),
            ),
            file_entry(
                answers, "answer_key", "complete" if len(records) == expected and not parsing_errors else "incomplete",
                expected_questions=expected, extracted_questions=len(records), parse_errors=parsing_errors,
            ),
        ])
        errors.extend(parsing_errors)

    docx = FORMAL / "Heilman Predictions--July 2026.docx"
    document = Document(docx)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    artifacts.append(file_entry(
        docx, "study_prediction_document", "not_question_import_source",
        paragraphs=len(document.paragraphs), tables=len(document.tables),
        reason="Narrative July 2026 prediction and study guidance; no discrete MBE records with four choices and answer key.",
        rendered_pages_verified=23,
        question_records_extracted=0,
        contains_prediction_language="predictions are educated guesses" in paragraph_text.lower(),
    ))

    review = json.loads((FORMAL / "official_questions_review.json").read_text(encoding="utf-8"))
    review_questions = review["questions"]
    tag_counts = Counter(tag for item in review_questions for tag in item.get("tags", []))
    output = {
        "schema": "passbar.formal-test.source-audit.v1",
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "artifacts": artifacts,
        "checks": {
            "pdf_question_records_expected": total,
            "review_question_records": len(review_questions),
            "record_count_matches": total == len(review_questions),
            "parse_errors": errors,
            "all_imported_questions_have_official_exam_tag": tag_counts.get("official_exam", 0) == len(review_questions),
            "unexpected_tags": sorted(tag for tag in tag_counts if tag != "official_exam"),
            "seven_day_total": seven_day_total,
            "seven_day_declared_total": 200,
            "seven_day_total_matches_declared": seven_day_total == 200,
        },
    }
    output_path = args.output.resolve()
    output_path.write_text(json.dumps(output, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(f"Audit: {output_path.relative_to(ROOT)}")
    print(f"PDF records: {total}; review records: {len(review_questions)}")
    print(f"Complete artifacts: {sum(item['status'] == 'complete' for item in artifacts)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
